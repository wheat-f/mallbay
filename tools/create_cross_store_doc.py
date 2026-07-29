from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION

OUT = r"D:\workSpace\mallbay\docs\features\cross-store-construction-solution.docx"

BLUE = "2E74B5"
DARK = "1F4D78"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
RED = "9B1C1C"

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(widths))); tblW.set(qn('w:type'), 'dxa')
    ind = tblPr.find(qn('w:tblInd'))
    if ind is None:
        ind = OxmlElement('w:tblInd'); tblPr.append(ind)
    ind.set(qn('w:w'), '120'); ind.set(qn('w:type'), 'dxa')
    grid = tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); grid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i] / 1440)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
            tcW.set(qn('w:w'), str(widths[i])); tcW.set(qn('w:type'), 'dxa')
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'), 'true'); trPr.append(tblHeader)

def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]; set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val)); r.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p

def add_number(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p

def add_callout(doc, title, text, fill=PALE, color=DARK):
    t = doc.add_table(rows=1, cols=1)
    set_table_geometry(t, [9360])
    c = t.cell(0, 0); set_cell_shading(c, fill)
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title + "  "); r.bold = True; r.font.color.rgb = RGBColor.from_string(color)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def configure_styles(doc):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'; normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    normal.font.size = Pt(11); normal.font.color.rgb = RGBColor(0x22,0x22,0x22)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ('Heading 1', 16, BLUE, 16, 8), ('Heading 2', 13, BLUE, 12, 6), ('Heading 3', 12, DARK, 8, 4)]:
        st = doc.styles[name]; st.font.name = 'Calibri'; st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after); st.paragraph_format.keep_with_next = True
    for s in ['List Bullet', 'List Bullet 2', 'List Number']:
        st = doc.styles[s]; st.font.name = 'Calibri'; st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei'); st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4); st.paragraph_format.line_spacing = 1.167

def build():
    doc = Document(); configure_styles(doc)
    # quiet running footer
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = footer.add_run('MallBay · 跨门店施工方案'); rr.font.size = Pt(9); rr.font.color.rgb = RGBColor(0x77,0x77,0x77)

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    r = p.add_run('跨门店外派施工业务方案'); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor.from_string(DARK)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(12)
    r = p.add_run('主门店建单 · 目标门店协同执行 · 统一留痕结算'); r.font.size = Pt(13); r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14)
    p.add_run('适用场景：').bold = True; p.add_run('施工人员通常归属 A 门店，但临时外派到 B/C 门店完成施工。本文用于业务确认、产品设计和实施验收。')
    add_callout(doc, '方案结论', 'A 门店创建并拥有销售订单；B/C 门店只接收并完成关联的外派施工任务。订单、批次、照片、工时和成本都通过同一施工任务关联，避免重复建单和数据断裂。', fill='E8F1FB')

    doc.add_heading('1. 业务目标与设计原则', level=1)
    for t in [
        '保留销售归属：客户、销售员、价格、收款、发票和销售业绩归订单主门店 A。',
        '明确执行归属：B/C 负责接收任务、排班、施工、照片、批次和异常处理。',
        '不重复建单：B/C 不创建第二张销售订单，只处理关联施工任务。',
        '执行和商业分离：B/C 默认不查看销售价格、毛利和收款信息。',
        '全过程可追溯：施工人员、批次、照片、工时、库存和跨店结算均保留审计记录。']:
        add_bullet(doc, t)

    doc.add_heading('2. 关键对象与归属关系', level=1)
    add_table(doc, ['对象', '归属/维护方', '说明'], [
        ('销售订单', 'A 门店', '客户、车辆、产品、销售价格、收款和发票的唯一商业载体。'),
        ('施工执行门店', 'A 选择；B/C确认', '订单创建时指定，外派场景必须填写。'),
        ('跨门店施工任务', '系统自动生成', '与销售订单一对一关联，承载排班、施工、照片、批次和异常。'),
        ('施工人员', '人员所属门店', '人员可归属 A，同时配置可外派到 B/C。'),
        ('材料批次', '实际出库门店', '必须记录来源门店、批次、数量、单位和成本。'),
        ('跨店结算单', '财务', '记录人工、材料、协作费和其他内部结算，不修改客户订单金额。')], [2100, 1800, 5460])

    doc.add_heading('3. 业务流程', level=1)
    for t in [
        'A 门店新建订单：施工方式选择“外派到其他门店”，填写 B/C、预约时间和施工要求。',
        '系统生成外派施工任务：任务带出订单、客户、车辆、产品、执行门店和预约信息。',
        'B/C 门店接收任务：确认施工容量；无法接收时必须填写拒绝原因并退回 A。',
        '派工与施工：A 店长指定外派师傅，B/C 施工主管可安排本地协助人员。',
        '施工执行：师傅领取或绑定批次，上传施工前/中/后照片，填写实际工时和异常。',
        'B/C 提交完成：系统校验批次、必需照片、工时和异常处理结果。',
        'A 门店最终验收：确认施工交付，销售订单继续进入交付或质保流程。',
        '财务结算：按实际材料、实际工时和门店协作费用形成跨店结算记录。']:
        add_number(doc, t)

    doc.add_heading('4. 任务状态设计', level=1)
    add_table(doc, ['状态', '责任方', '可执行操作'], [
        ('待目标门店确认', 'B/C 店长', '接收、拒绝、申请调整时间。'),
        ('已接收待派工', 'A 店长/B/C主管', '指定外派师傅或安排协助人员。'),
        ('已派工', '施工人员', '查看任务、准备材料和到场施工。'),
        ('施工中', '施工人员', '绑定批次、上传过程照片、记录工时。'),
        ('待补充证据', '施工人员/B/C主管', '补齐照片、批次或异常说明。'),
        ('待主门店验收', 'A 店长', '查看资料、退回补充或确认完成。'),
        ('已完成', '系统', '冻结施工证据并进入订单交付/质保流程。')], [2400, 1800, 5160])
    add_callout(doc, '重要约束', 'B/C 只能完成施工任务，不能直接完成销售订单；销售订单的交付和最终商业状态仍由 A 门店确认。', fill='FFF4D6', color='7A5A00')

    doc.add_heading('5. B/C 门店功能设计', level=1)
    add_bullet(doc, '新增“外派施工任务”菜单，按待接收、待派工、施工中、待验收、已完成分类。')
    add_bullet(doc, '任务详情展示客户、车辆、施工项目、预约时间、来源门店、人员、材料和注意事项。')
    add_bullet(doc, '默认隐藏销售单价、收款金额、欠款、毛利和销售提成。')
    add_bullet(doc, 'B/C 可上传施工资料、确认现场状态、安排协助人员和反馈异常。')
    add_bullet(doc, '不得修改客户、产品金额、销售员、收款、发票或销售订单状态。')

    doc.add_heading('6. 人员、照片与工时', level=1)
    add_table(doc, ['记录项', '最低要求', '验收规则'], [
        ('人员', '记录所属门店、施工角色和任务分工', '每个实际参与人员必须有任务成员记录。'),
        ('施工前照片', '外观/施工区域清晰可辨', '缺失时不能提交施工完成。'),
        ('施工中照片', '关键工序或材料使用记录', '异常任务必须补充。'),
        ('施工后照片', '完成效果和交付状态', '至少一张，支持客户确认。'),
        ('实际工时', '每人填写开始、结束或确认分钟数', '未确认前只统计实际提成，不虚构成本分摊。')], [2000, 3400, 3960])
    add_callout(doc, '照片统一归档', '所有照片绑定订单号和施工任务号，记录上传人、门店、时间和照片类型。B/C 上传后 A 可查看，不需要重复上传。')

    doc.add_heading('7. 库存、批次和成本', level=1)
    doc.add_paragraph('跨门店施工必须明确材料来源，不能只记录一个数量。推荐优先使用执行门店实际领料，并将出库批次直接绑定到施工任务。')
    add_table(doc, ['场景', '库存处理', '成本归集'], [
        ('B/C 使用本店库存', 'B/C 按批次出库并绑定任务', '按 B/C 实际出库价回传 A 订单。'),
        ('A 带料外派', 'A 出库或调拨至 B/C，再由任务消耗', '按 A 实际出库/调拨成本归集。'),
        ('B/C 临时采购', '采购入库后绑定任务批次', '按实际入库价计入订单成本。'),
        ('零散产品施工', '支持米、卷等销售/库存单位换算', '保留原批次和换算关系，不使用手工库存调整代替。')], [2000, 3800, 3560])

    doc.add_heading('8. 跨门店结算', level=1)
    for t in [
        '销售收入和客户收款仍归 A 门店。',
        '外派人员人工成本按每人确认的实际工时分摊。',
        'B/C 协助人员的人工成本按实际工时单独记录。',
        '材料成本按实际入库价优先，缺失时使用产品材料成本标准兜底，并标记成本来源。',
        '门店协作费、交通费等作为内部结算项目，不直接改变客户订单收费。']:
        add_bullet(doc, t)
    add_table(doc, ['结算主体', '适用情况', '处理方式'], [
        ('同一法人门店', 'A、B、C 属于同一财务主体', '形成内部门店结算，记录应付方、应收方和审批。'),
        ('不同法人门店', 'B/C 为外部合作主体', '按外包/供应商流程形成费用、付款和凭证。'),
        ('财务', '统一处理结算', '只处理成本和付款，不修改销售订单业务字段。')], [2200, 3100, 4060])

    doc.add_heading('9. 权限矩阵', level=1)
    add_table(doc, ['角色', '允许操作', '禁止操作'], [
        ('A 店长', '建单、选择执行门店、派工、验收、查看全部信息', '无'),
        ('B/C 店长', '接收任务、排班、安排协助、查看施工资料', '修改销售金额、收款、发票、销售员'),
        ('施工主管', '排班、派工、审核施工证据、处理异常', '修改客户订单商业字段'),
        ('施工师傅', '查看任务、领料、上传照片、确认工时', '修改订单金额和收款'),
        ('财务', '查看成本、跨店结算和付款', '派工、修改订单、改变施工状态'),
        ('总部管理员', '跨门店查看、审计和配置', '绕过审批直接删除已确认记录')], [1800, 4100, 3460])

    doc.add_heading('10. 异常与通知', level=1)
    add_bullet(doc, 'B/C 拒绝接收：必须填写原因，A 门店收到通知后重新选择门店或调整预约。')
    add_bullet(doc, '外派师傅无法到场：B/C 发起换人请求，不能直接虚构施工完成。')
    add_bullet(doc, '材料不足：选择调拨、使用本店库存、临时采购或改期，形成库存流水。')
    add_bullet(doc, '照片/工时不完整：任务进入“待补充证据”，不能提交施工完成。')
    add_bullet(doc, '预约临近未接收：通知 B/C 店长和 A 店长；超时自动升级提醒。')

    doc.add_heading('11. 分期实施计划', level=1)
    add_table(doc, ['阶段', '范围', '验收结果'], [
        ('一期：任务协同', '执行门店、任务生成、接收/拒绝、派工、施工状态、照片共享、权限隔离', 'B/C 能看到并完成外派任务，A 能查看全过程。'),
        ('二期：库存批次', '跨店领料、调拨、批次绑定、零散出库、实际成本回传', '每次出库都能追溯到订单和施工任务。'),
        ('三期：工时结算', '实际工时、人员成本、协作费、跨店结算和财务审批', '成本和费用按实际发生结算。'),
        ('四期：分析报表', '外派任务量、完成率、工时、材料、成本、延期和售后率', '支持门店和人员绩效分析。')], [1700, 4800, 2860])

    doc.add_heading('12. 建议的数据对象', level=1)
    doc.add_paragraph('建议新增“跨门店施工任务”及其关联对象，不新增第二张销售订单。')
    add_table(doc, ['对象', '关键字段'], [
        ('CrossStoreExecutionTask', 'taskId、orderId、sourceStoreId、executionStoreId、status、预约时间、拒绝原因、结算状态'),
        ('TaskMember', 'taskId、userId、storeId、角色、计划工时、实际工时、实际成本、提成'),
        ('TaskMaterial', 'taskId、productId、batchId、来源门店、数量、单位、实际成本、库存流水号'),
        ('TaskEvidence', 'taskId、照片类型、文件地址、上传人、上传门店、上传时间'),
        ('CrossStoreSettlement', 'taskId、付款门店、收款门店、人工、材料、协作费、审批人、结算时间')], [2500, 6860])

    doc.add_heading('13. 待确认的业务参数', level=1)
    add_callout(doc, '实施前确认', '以下参数不影响总体架构，但会影响库存和财务结算实现：A/B/C 是否属于同一法人或财务主体；外派施工默认由 A 提供材料，还是由执行门店提供材料；B/C 是否可以收取固定协作服务费。', fill='FFF4D6', color='7A5A00')
    doc.add_paragraph('默认建议：同一法人采用内部门店结算；材料优先由执行门店实际领料；所有成本和协作费单独结算，不改变客户订单金额。')

    doc.add_heading('14. 验收清单', level=1)
    for t in [
        'A 创建外派订单后，B/C 能收到且不产生重复销售订单。',
        'B/C 可以接收、拒绝、派工、施工和提交完成。',
        'A 可以查看施工照片、批次、实际工时和异常记录。',
        '施工完成前，缺少必需照片或批次时系统能够阻止提交。',
        '库存出库、零散换算和成本来源均可追溯到订单和施工任务。',
        'B/C 和施工人员不能修改销售价格、收款和发票。',
        '财务可以查看跨店成本和结算，但不能派工或修改订单。']:
        add_bullet(doc, '□ ' + t)

    doc.save(OUT)
    print(OUT)

if __name__ == '__main__': build()
